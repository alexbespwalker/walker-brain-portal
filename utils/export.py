"""Export utilities for CSV and Word document generation."""

import io
import pandas as pd
import streamlit as st


def download_csv(df: pd.DataFrame, filename: str = "export.csv") -> None:
    """Render a right-aligned CSV download button for the given DataFrame."""
    csv_buf = io.StringIO()
    df.to_csv(csv_buf, index=False)
    _, btn_col = st.columns([3, 1])
    with btn_col:
        st.download_button(
            label="Export CSV",
            data=csv_buf.getvalue(),
            file_name=filename,
            mime="text/csv",
            type="primary",
        )


def format_quote_for_clipboard(
    quote: str,
    case_type: str | None = None,
    tone: str | None = None,
    quality: int | None = None,
    date: str | None = None,
) -> str:
    """Format a quote with context for creative briefs."""
    parts = [f'"{quote}"']
    meta = []
    if case_type:
        meta.append(f"Case type: {case_type}")
    if tone:
        meta.append(f"Tone: {tone}")
    if quality is not None:
        meta.append(f"Quality: {quality}/100")
    if date:
        meta.append(f"Date: {date}")
    if meta:
        parts.append(" | ".join(meta))
    return "\n".join(parts)


def generate_word_doc(title: str, content_blocks: list[dict]) -> bytes:
    """Generate a Word document from content blocks.

    Each block: {"heading": str, "body": str}
    """
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)

    for block in content_blocks:
        if block.get("heading"):
            doc.add_heading(block["heading"], level=1)
        if block.get("body"):
            doc.add_paragraph(block["body"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
